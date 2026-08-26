"""
JARVIS plugin — turns a PDF, DOCX, or TXT file into structured CSV/JSON/XLSX,
or a cleaned-up Word document, using an LLM to make sense of messy content.

PDF text comes from PyMuPDF (pymupdf) page by page; any page with no real
text layer (a scan or a phone photo) automatically falls back to local
Tesseract OCR via core.ocr_utils, so a scanned PDF works here too — no
pdfplumber dependency, no separate step needed before calling this plugin.

For well-structured documents where exact table values matter and no AI
massaging is wanted, use document_processor instead. For OCR text only (no
structuring into a document), use document_ocr.
"""

import csv
import json
from pathlib import Path

try:
    import docx
except Exception:
    docx = None

try:
    import pandas as pd
except Exception:
    pd = None

from core.ai_text import generate as get_ai_response

PLUGIN = {
    "name": "document_extractor",
    "description": (
        "Extracts data from a PDF, DOCX, or TXT file — including a SCANNED or "
        "PHOTOGRAPHED PDF, via automatic OCR fallback — and converts it into "
        "CSV, JSON, XLSX, or a cleaned-up Word (DOCX) document using an AI "
        "prompt. Use for: 'bu PDF'i excele çevir', 'bu taranmış faturayı "
        "excele/word'e dönüştür', 'bu belgeyi düzenli bir word dosyası yap'."
    ),
    "parameters": {
        "type": "OBJECT",
        "properties": {
            "file_path": {
                "type": "STRING",
                "description": "Absolute or relative path to the source document (txt, pdf, or docx)."
            },
            "output_format": {
                "type": "STRING",
                "enum": ["csv", "json", "xlsx", "docx"],
                "description": "Desired output format for the extracted data."
            },
            "prompt": {
                "type": "STRING",
                "description": "Optional custom prompt to guide the AI on how to structure the extracted data."
            }
        },
        "required": ["file_path", "output_format"]
    }
}


def _extract_pdf(path: Path, log=None):
    from core.ocr_utils import ensure_tesseract_ready, extract_pdf_pages, extract_pdf_tables

    ready, msg = ensure_tesseract_ready()
    if ready:
        pages = extract_pdf_pages(path, log=log)
    else:
        # Text-layer pages still work without Tesseract; only OCR pages would be blank.
        import pymupdf as fitz
        doc = fitz.open(path)
        pages = [page.get_text("text") for page in doc]
        doc.close()
        if log:
            log(f"Uyarı: {msg} Taranmış sayfalar atlanacak.")

    text = "\n".join(pages)
    tables = extract_pdf_tables(path)
    return text, tables


def _extract_text_from_docx(path: Path) -> str:
    if not docx:
        raise RuntimeError("python-docx library is not installed.")
    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs)


def _extract_tables_from_docx(path: Path):
    if not docx:
        raise RuntimeError("python-docx library is not installed.")
    doc = docx.Document(str(path))
    tables = []
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        tables.append(data)
    return tables


def _strip_code_fence(text: str) -> str:
    """Drop a leading/trailing ``` fence, e.g. ```csv\n...\n``` -- some models
    wrap their answer in one despite being told not to."""
    stripped = text.strip()
    if stripped.startswith("```"):
        lines = stripped.splitlines()
        lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        stripped = "\n".join(lines)
    return stripped


def _write_tabular_output(data, fmt: str, base_name: str) -> str:
    output_path = Path(f"{base_name}.{fmt}")
    if fmt == "json":
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    elif fmt == "csv":
        with open(output_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            for row in data:
                writer.writerow(row)
    elif fmt == "xlsx":
        if not pd:
            raise RuntimeError("pandas is required for XLSX output but is not installed.")
        df = pd.DataFrame(data[1:], columns=data[0] if data else None)
        df.to_excel(output_path, index=False)
    else:
        raise ValueError(f"Unsupported format: {fmt}")
    return str(output_path)


def _write_docx_output(cleaned_text: str, raw_tables: list, base_name: str) -> str:
    if not docx:
        raise RuntimeError("python-docx library is not installed.")
    output_path = Path(f"{base_name}.docx")
    document = docx.Document()
    for para in cleaned_text.split("\n\n"):
        para = para.strip()
        if para:
            document.add_paragraph(para)

    for idx, table_rows in enumerate(raw_tables, start=1):
        if not table_rows:
            continue
        document.add_heading(f"Tablo {idx}", level=2)
        n_cols = max(len(r) for r in table_rows)
        word_table = document.add_table(rows=len(table_rows), cols=n_cols)
        try:
            word_table.style = "Light Grid Accent 1"
        except Exception:
            pass
        for r_idx, row in enumerate(table_rows):
            for c_idx in range(n_cols):
                cell_text = row[c_idx] if c_idx < len(row) and row[c_idx] is not None else ""
                word_table.cell(r_idx, c_idx).text = str(cell_text)

    document.save(output_path)
    return str(output_path)


def run(parameters: dict, player=None, session_memory=None) -> str:
    message, _output_path = _run_impl(parameters, player, session_memory)
    return message


def run_and_get_path(parameters: dict, player=None, session_memory=None) -> tuple[str, str | None]:
    """Same as run(), but also returns the produced file's path (or None on
    failure) -- for callers that need the actual file, not just the spoken
    confirmation text (e.g. telegram_bridge sending it back as an attachment).
    """
    return _run_impl(parameters, player, session_memory)


def _run_impl(parameters: dict, player=None, session_memory=None) -> tuple[str, str | None]:
    try:
        file_path = parameters.get("file_path")
        output_format = (parameters.get("output_format") or "").lower()
        custom_prompt = parameters.get("prompt")

        if not file_path:
            return "Belirttiğiniz dosya yolu eksik. Lütfen bir dosya yolu sağlayın.", None
        if output_format not in {"csv", "json", "xlsx", "docx"}:
            return f"Desteklenmeyen çıktı formatı: {output_format}. Lütfen csv, json, xlsx veya docx seçin.", None

        path = Path(file_path).expanduser().resolve()
        if not path.is_file():
            return f"Dosya bulunamadı: {file_path}. Lütfen yolu kontrol edin.", None

        def log(m: str) -> None:
            if player:
                try:
                    player.write_log(f"JARVIS: {m}")
                except Exception:
                    pass

        ext = path.suffix.lower()
        raw_text = ""
        tables = []
        if ext == ".txt":
            raw_text = path.read_text(encoding="utf-8")
        elif ext == ".pdf":
            log(f"'{path.name}' okunuyor (taranmış sayfalar varsa OCR devreye girer)…")
            raw_text, tables = _extract_pdf(path, log=log)
        elif ext == ".docx":
            raw_text = _extract_text_from_docx(path)
            tables = _extract_tables_from_docx(path)
        else:
            return f"Dosya uzantısı {ext} desteklenmiyor. Sadece txt, pdf ve docx dosyaları işlenebilir.", None

        if not raw_text.strip() and not tables:
            return f"Sir, '{path.name}' içinden okunabilir içerik çıkaramadım.", None

        base_name = path.stem + "_extracted"

        if output_format == "docx":
            prompt = custom_prompt or (
                "Aşağıdaki belge metnini oku ve düzgün, okunabilir paragraflar "
                "halinde (gereksiz OCR gürültüsünü, tekrar eden başlıkları "
                "temizleyerek) yeniden yaz. Sadece temizlenmiş metni döndür, "
                "ek açıklama ekleme."
            )
            cleaned_text = get_ai_response(f"{prompt}\n\n---\n\n{raw_text}")
            output_file = _write_docx_output(cleaned_text, tables, base_name)
            return f"Belge başarıyla işlendi ve Word dosyası olarak kaydedildi: {output_file}.", output_file

        if custom_prompt:
            prompt = custom_prompt
        else:
            prompt = (
                "Aşağıdaki belge içeriğini ve varsa tabloları analiz et ve "
                "verileri düz metin CSV formatına dönüştür: her satır bir "
                "veri satırı, sütunlar virgülle ayrılmış, ilk satır başlıklar. "
                "SADECE ham CSV satırlarını döndür — kod bloğu, açıklama, "
                "başka bir dosya formatı (xlsx/base64/binary) veya markdown "
                "kullanma."
            )
        content_for_ai = raw_text
        if tables:
            content_for_ai += "\n\nTablolar:\n"
            for idx, tbl in enumerate(tables, start=1):
                content_for_ai += f"Tablo {idx}:\n"
                for row in tbl:
                    content_for_ai += ", ".join(str(c) if c is not None else "" for c in row) + "\n"
                content_for_ai += "\n"

        ai_response = get_ai_response(f"{prompt}\n\n---\n\n{content_for_ai}")

        if output_format == "json":
            try:
                structured_data = json.loads(_strip_code_fence(ai_response))
            except Exception:
                structured_data = {"result": ai_response}
        else:  # csv, xlsx
            cleaned = _strip_code_fence(ai_response)
            structured_data = [line.split(",") for line in cleaned.strip().splitlines() if line.strip()]

        output_file = _write_tabular_output(structured_data, output_format, base_name)
        return f"Belge başarıyla işlendi ve {output_format.upper()} dosyası olarak kaydedildi: {output_file}.", output_file
    except Exception as exc:
        return f"Belge işleme sırasında bir hata oluştu: {str(exc)}", None
